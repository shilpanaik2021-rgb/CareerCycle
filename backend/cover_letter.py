"""
AI-powered cover letter generator using Mistral AI.
Generates personalized cover letters based on Shilpa's resume and job descriptions.
Saves cover letters as .txt files in /backend/cover_letters/.
"""
import os
import re
import json
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import resume_data
import config

# Directory for cover letter text files
COVER_LETTERS_DIR = os.path.join(os.path.dirname(__file__), "cover_letters")
os.makedirs(COVER_LETTERS_DIR, exist_ok=True)


def _get_mistral_api_key():
    """Retrieve Mistral API Key."""
    if not config.MISTRAL_API_KEY:
        raise ValueError("Mistral API key not configured. Add MISTRAL_API_KEY to .env file.")
    return config.MISTRAL_API_KEY


def _build_prompt(job_title: str, company: str, job_description: str) -> str:
    """Build the cover letter generation prompt with Shilpa's style guide."""
    resume_text = resume_data.get_full_resume_text()
    style = resume_data.COVER_LETTER_STYLE

    prompt = f"""Write a professional cover letter for the following job posting.
The cover letter must be written in the voice and style described below.

═══ APPLICANT'S RESUME ═══
{resume_text}

═══ JOB DETAILS ═══
Position: {job_title}
Company: {company}
Job Description:
{job_description[:3000]}

═══ WRITING STYLE REQUIREMENTS ═══
1. Opening: Start with "{style['opening'].format(role=job_title, company=company)}"
2. Use these confidence phrases naturally throughout: {', '.join(style['confidence_phrases'])}
3. Highlight these key qualifications where relevant to the job:
   {chr(10).join('   - ' + h for h in style['key_highlights'])}
4. Reference specific requirements from the job posting and match them to the applicant's experience
5. If the job mentions Epic, ECW, or any EHR system, emphasize the applicant's Epic Superuser status
6. If the job mentions budget or financial management, reference the $1.6M budget management experience
7. If the job mentions team leadership, reference managing 25+ billing specialists
8. Mention the MHA degree as a differentiator
9. Closing: End with "{style['closing']}"
10. Sign off: "{style['signature']}"

═══ FORMAT REQUIREMENTS ═══
- Write 3-4 body paragraphs (not including opening and closing)
- Total length: 350-450 words
- Professional but warm tone
- Do NOT include a date or address header — just the letter body
- Do NOT include "Dear Hiring Manager" — start directly with the opening line
- End with the exact signature block provided above

Write the complete cover letter now:"""

    return prompt


def generate_cover_letter(job_id: str, job_title: str, company: str, job_description: str) -> str:
    """
    Generate a cover letter using Mistral AI.
    Returns the cover letter text and saves it as a .txt file.
    """
    mistral_key = _get_mistral_api_key()
    prompt = _build_prompt(job_title, company, job_description)

    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Authorization": f"Bearer {mistral_key}"
    }
    payload = {
        "model": "mistral-large-latest",
        "messages": [{"role": "user", "content": prompt}]
    }

    try:
        response = requests.post("https://api.mistral.ai/v1/chat/completions", headers=headers, json=payload, timeout=30)
        if response.status_code != 200:
            raise ValueError(f"Mistral API error: {response.text}")
        data = response.json()
        letter_text = data['choices'][0]['message']['content'].strip()
    except Exception as e:
        raise ValueError(f"Failed to generate cover letter: {str(e)}")

    # Save as .txt file
    safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')[:30]
    safe_title = re.sub(r'[^\w\s-]', '', job_title).strip().replace(' ', '_')[:30]
    filename = f"cover_letter_{job_id}_{safe_company}_{safe_title}.txt"
    filepath = os.path.join(COVER_LETTERS_DIR, filename)

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(letter_text)

    return letter_text, filepath


def generate_cover_letter_docx(job_id: str, job_title: str, company: str, letter_text: str) -> str:
    """
    Generate a .docx file from cover letter text.
    Returns the path to the generated file.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Add contact header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run(f"{resume_data.PERSONAL_INFO['full_name_with_suffix']}\n")
    header_run.bold = True
    header_run.font.size = Pt(14)

    contact_run = header.add_run(
        f"{resume_data.PERSONAL_INFO['location']} | "
        f"{resume_data.PERSONAL_INFO['phone']} | "
        f"{resume_data.PERSONAL_INFO['email']}\n"
        f"LinkedIn: {resume_data.PERSONAL_INFO['linkedin']}"
    )
    contact_run.font.size = Pt(10)

    # Add a line break
    doc.add_paragraph()

    # Add the cover letter body
    paragraphs = letter_text.split("\n\n")
    for para_text in paragraphs:
        if para_text.strip():
            p = doc.add_paragraph(para_text.strip())
            p.paragraph_format.space_after = Pt(8)

    # Save the document
    safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')[:30]
    filename = f"Cover_Letter_{safe_company}_{job_id}.docx"
    filepath = os.path.join(COVER_LETTERS_DIR, filename)
    doc.save(filepath)

    return filepath


def generate_all_missing(jobs_df, log_list: list) -> int:
    """
    Generate cover letters for all jobs that don't have one using Mistral.
    Returns the count of letters generated.
    """
    import pandas as pd

    try:
        mistral_key = _get_mistral_api_key()
    except Exception as e:
        log_list.append({"type": "error", "message": f"❌ Mistral API Configuration Error: {str(e)}"})
        return 0

    count = 0
    missing = jobs_df[
        (jobs_df["cover_letter_path"].isna()) | (jobs_df["cover_letter_path"] == "")
    ]

    log_list.append({"type": "progress", "message": f"⚡ Generating cover letters for {len(missing)} jobs..."})

    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
        "Authorization": f"Bearer {mistral_key}"
    }

    for idx, row in missing.iterrows():
        job_id = row["id"]
        title = row["title"]
        company = row["company"]
        description = str(row.get("description", ""))

        if not description or description == "nan":
            log_list.append({"type": "info", "message": f"   ⏭️ Skipping {company} - {title} (no job description)"})
            continue

        try:
            log_list.append({"type": "progress", "message": f"⚡ [{count+1}] Generating for {company} - {title}..."})

            prompt = _build_prompt(title, company, description)
            payload = {
                "model": "mistral-large-latest",
                "messages": [{"role": "user", "content": prompt}]
            }

            response = requests.post("https://api.mistral.ai/v1/chat/completions", headers=headers, json=payload, timeout=30)
            if response.status_code != 200:
                raise ValueError(f"Mistral API error: {response.text}")
            
            data = response.json()
            letter_text = data['choices'][0]['message']['content'].strip()

            # Save as .txt file
            safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')[:30]
            safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')[:30]
            filename = f"cover_letter_{job_id}_{safe_company}_{safe_title}.txt"
            filepath = os.path.join(COVER_LETTERS_DIR, filename)

            with open(filepath, "w", encoding="utf-8") as f:
                f.write(letter_text)

            # Update the DataFrame
            jobs_df.at[idx, "cover_letter_path"] = filepath
            count += 1

            log_list.append({"type": "success", "message": f"   ✅ Generated cover letter for {company}"})

        except Exception as e:
            log_list.append({"type": "error", "message": f"   ❌ Error for {company}: {str(e)}"})

    # Save updated DataFrame
    jobs_df.to_csv(config.JOBS_CSV_PATH, index=False)

    log_list.append({"type": "success", "message": f"✅ Cover letter generation complete! {count} letters created."})
    return count
