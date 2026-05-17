import os
import re
import json
import docx
import requests
import urllib.parse
from bs4 import BeautifulSoup
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyPDF2 import PdfReader

# ReportLab Imports
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

# Ensure uploads directory exists
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "resume_uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)
RESUME_TEXT_PATH = os.path.join(UPLOAD_DIR, "resume_text.txt")
IMPROVED_RESUME_PATH = os.path.join(UPLOAD_DIR, "improved_resume.txt")

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extracts all text from a PDF file using PyPDF2."""
    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text.strip()

def get_mistral_api_key():
    """Initializes and returns the Mistral API Key."""
    api_key = os.getenv("MISTRAL_API_KEY")
    if not api_key:
        raise ValueError("MISTRAL_API_KEY is not configured in settings.")
    return api_key

def free_web_search(query: str, max_results=3):
    """Performs a free, zero-config DuckDuckGo web search via the lite HTML site."""
    url = "https://lite.duckduckgo.com/lite/"
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }
    try:
        response = requests.post(url, headers=headers, data={"q": query}, timeout=5)
        if response.status_code != 200:
            return []
        
        soup = BeautifulSoup(response.text, 'html.parser')
        results = []
        for link in soup.find_all('a', class_='result-link')[:max_results]:
            title = link.get_text()
            href = link.get('href')
            results.append({
                "title": title.strip(),
                "url": href
            })
        return results
    except Exception as e:
        print("Free web search error:", e)
        return []

def parse_analysis_scores(text: str):
    """Parses score and breakdown from response using regex."""
    overall_match = re.search(r'OVERALL SCORE:\s*\[?(\d+)\]?/100', text, re.IGNORECASE)
    overall = int(overall_match.group(1)) if overall_match else 75
    
    breakdown = {
        "ats": 15,
        "achievements": 15,
        "keywords": 15,
        "formatting": 15,
        "narrative": 15
    }
    
    ats_match = re.search(r'ATS Compatibility:\s*(\d+)/20', text, re.IGNORECASE)
    if ats_match: breakdown["ats"] = int(ats_match.group(1))
    
    ach_match = re.search(r'Quantified Achievements:\s*(\d+)/20', text, re.IGNORECASE)
    if ach_match: breakdown["achievements"] = int(ach_match.group(1))
    
    key_match = re.search(r'Keywords & Skills:\s*(\d+)/20', text, re.IGNORECASE)
    if key_match: breakdown["keywords"] = int(key_match.group(1))
    
    fmt_match = re.search(r'Formatting & Clarity:\s*(\d+)/20', text, re.IGNORECASE)
    if fmt_match: breakdown["formatting"] = int(fmt_match.group(1))
    
    nar_match = re.search(r'Career Narrative:\s*(\d+)/20', text, re.IGNORECASE)
    if nar_match: breakdown["narrative"] = int(nar_match.group(1))
    
    return overall, breakdown

async def stream_gemini_analysis(resume_text: str):
    """Streams resume analysis with free search grounding and Mistral Large model."""
    try:
        mistral_key = get_mistral_api_key()
    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
        return

    # 1. Perform 2 grounding searches
    queries = [
        "Orlando FL medical billing manager salary trends 2025",
        "Epic EHR medical billing manager resume keywords 2025"
    ]
    
    grounding_data = []
    
    for q in queries:
        yield f"data: {json.dumps({'type': 'search_start', 'query': q})}\n\n"
        search_results = free_web_search(q, max_results=3)
        for res in search_results:
            yield f"data: {json.dumps({'type': 'search_result', 'url': res['url'], 'title': res['title']})}\n\n"
            grounding_data.append(res)
            
    # Compile grounding context for Mistral
    grounding_str = ""
    if grounding_data:
        grounding_str = "\n═══ REAL-TIME WEB SEARCH RESULTS (GROUNDING) ═══\n"
        for idx, res in enumerate(grounding_data):
            grounding_str += f"[{idx+1}] Title: {res['title']}\n    URL: {res['url']}\n"

    system_prompt = (
        "You are an expert resume coach and healthcare industry recruiter with "
        "15 years of experience hiring for revenue cycle management, medical billing, "
        "and healthcare administration roles. You have deep knowledge of what "
        "healthcare employers look for in 2025.\n\n"
        "The user has uploaded their resume. Your job is to:\n"
        "1. Analyze every section of this resume critically and honestly\n"
        "2. Incorporate real-time web search insights from the search results provided below\n"
        "3. Give specific, actionable feedback — not generic advice\n\n"
        "Structure your response EXACTLY like this:\n\n"
        "## OVERALL SCORE: [X]/100\n\n"
        "## SCORE BREAKDOWN:\n"
        "- ATS Compatibility: X/20 — [one line explanation]\n"
        "- Quantified Achievements: X/20 — [one line explanation]\n"
        "- Keywords & Skills: X/20 — [one line explanation]\n"
        "- Formatting & Clarity: X/20 — [one line explanation]\n"
        "- Career Narrative: X/20 — [one line explanation]\n\n"
        "## WHAT'S WORKING WELL ✅\n"
        "[3-5 specific genuine strengths with quotes from their resume]\n\n"
        "## CRITICAL ISSUES TO FIX ❌\n"
        "For each issue:\n"
        "### Issue [N]: [Issue Title]\n"
        "**Section:** [which part of resume]\n"
        "**Problem:** [specific explanation of why this is weak]\n"
        "**Current text:** \"[exact quote from their resume]\"\n"
        "**Why it hurts you:** [impact on hiring chances]\n"
        "**Fix it like this:** \"[rewritten version]\"\n\n"
        "## MISSING KEYWORDS 🔍\n"
        "[List of keywords employers search for that are NOT in their resume,\n"
        "based on your web research of current job postings]\n\n"
        "## ACTION PLAN 📋\n"
        "Priority 1 (do today): ...\n"
        "Priority 2 (this week): ...\n"
        "Priority 3 (nice to have): ...\n\n"
        "## MARKET INSIGHTS 🌐\n"
        "[Based on real-time web search results: what the current market is paying, "
        "what skills are trending, what hiring managers say they want]\n\n"
        f"Real-Time Grounding Context:{grounding_str}\n\n"
        f"Resume Text:\n{resume_text}"
    )

    try:
        headers = {
            "Content-Type": "application/json",
            "Accept": "application/json",
            "Authorization": f"Bearer {mistral_key}"
        }
        payload = {
            "model": "mistral-large-latest",
            "messages": [{"role": "user", "content": system_prompt}],
            "stream": True
        }
        
        response = requests.post("https://api.mistral.ai/v1/chat/completions", headers=headers, json=payload, stream=True, timeout=30)
        if response.status_code != 200:
            raise ValueError(f"Mistral API returned status code {response.status_code}: {response.text}")

        full_text = ""
        for line in response.iter_lines():
            if line:
                decoded = line.decode('utf-8').strip()
                if decoded.startswith("data: "):
                    data_str = decoded[6:]
                    if data_str == "[DONE]":
                        break
                    try:
                        data_json = json.loads(data_str)
                        content = data_json['choices'][0]['delta'].get('content', '')
                        if content:
                            full_text += content
                            yield f"data: {json.dumps({'type': 'text', 'content': content})}\n\n"
                    except:
                        pass
        
        overall, breakdown = parse_analysis_scores(full_text)
        yield f"data: {json.dumps({'type': 'score', 'overall': overall, 'breakdown': breakdown})}\n\n"
        yield f"data: {json.dumps({'type': 'done'})}\n\n"
    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'message': f"Mistral API Error: {str(e)}"})}\n\n"

async def stream_gemini_improvement(resume_text: str, suggestions: str = ""):
    """Streams full rewritten and optimized resume section by section using Mistral."""
    try:
        mistral_key = get_mistral_api_key()
    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
        return

    prompt = (
        "You are an expert resume writer specializing in high-end healthcare administration and revenue cycle management resumes.\n"
        "The user has requested to improve their resume based on recent ATS analysis.\n"
        f"Specific improvements suggested: {suggestions}\n\n"
        "Take the original resume and rewrite the entire content to be perfect. Make sure it incorporates:\n"
        "- Impactful action verbs\n"
        "- Quantified results ($ saved, % collection rates improved, denials reduced)\n"
        "- Crucial healthcare billing keywords (Epic EHR, medical coding, HIPAA, RCM, etc.)\n\n"
        "═══ STRICTION FORMATTING RULES ═══\n"
        "1. DO NOT include any contact information header at the top (name, address, email, phone, LinkedIn, etc.) — this is added automatically by the PDF template. Start directly with the `# PROFESSIONAL SUMMARY` section.\n"
        "2. Do NOT use markdown bold asterisks `**` or italic asterisks `*` inside the resume section headings or job headers. Keep all section headings perfectly clean plain text.\n"
        "3. Use exactly `# ` for main section headers (e.g., `# PROFESSIONAL SUMMARY`, `# PROFESSIONAL EXPERIENCE`, `# EDUCATION`, `# SKILLS & CERTIFICATIONS`).\n"
        "4. Use exactly `## ` for job title/company or degree lines (e.g., `## REVENUE CYCLE MANAGER | ADVENTHEALTH | ORLANDO, FL`).\n"
        "5. Stream back ONLY the clean, professional, fully polished resume sections without any introductory/explanatory notes or markdown wrappers.\n\n"
        f"Original Resume:\n{resume_text}"
    )

    try:
        headers = {
            "Content-Type": "application/json",
            "Accept": "application/json",
            "Authorization": f"Bearer {mistral_key}"
        }
        payload = {
            "model": "mistral-large-latest",
            "messages": [{"role": "user", "content": prompt}],
            "stream": True
        }
        
        response = requests.post("https://api.mistral.ai/v1/chat/completions", headers=headers, json=payload, stream=True, timeout=30)
        if response.status_code != 200:
            raise ValueError(f"Mistral API returned status code {response.status_code}")

        full_text = ""
        for line in response.iter_lines():
            if line:
                decoded = line.decode('utf-8').strip()
                if decoded.startswith("data: "):
                    data_str = decoded[6:]
                    if data_str == "[DONE]":
                        break
                    try:
                        data_json = json.loads(data_str)
                        content = data_json['choices'][0]['delta'].get('content', '')
                        if content:
                            full_text += content
                            yield f"data: {json.dumps({'type': 'text', 'content': content})}\n\n"
                    except:
                        pass
                        
        with open(IMPROVED_RESUME_PATH, "w", encoding="utf-8") as f:
            f.write(full_text)
            
        yield f"data: {json.dumps({'type': 'done'})}\n\n"
    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'message': f"Mistral API Error: {str(e)}"})}\n\n"

async def stream_gemini_chat(message: str, history: list, resume_text: str):
    """Streams live chat response with resume context, history, and real-time grounding using Mistral."""
    try:
        mistral_key = get_mistral_api_key()
    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
        return

    # Optional: run search grounding for Mistral chat if relevant
    grounding_data = []
    if any(k in message.lower() for k in ["salary", "job", "market", "pay", "trend", "hiring"]):
        yield f"data: {json.dumps({'type': 'search_start', 'query': message[:50]})}\n\n"
        search_results = free_web_search(message[:50], max_results=3)
        for res in search_results:
            yield f"data: {json.dumps({'type': 'search_result', 'url': res['url'], 'title': res['title']})}\n\n"
            grounding_data.append(res)

    grounding_str = ""
    if grounding_data:
        grounding_str = "\n═══ REAL-TIME CHAT GROUNDING RESULTS ═══\n"
        for idx, res in enumerate(grounding_data):
            grounding_str += f"[{idx+1}] Title: {res['title']}\n    URL: {res['url']}\n"

    formatted_history = ""
    for h in history:
        sender = "User" if h.get("sender") == "user" else "Assistant"
        formatted_history += f"{sender}: {h.get('text')}\n"
        
    prompt = (
        "You are a stellar career coach. You have full context of the user's resume.\n"
        "Answer the user's question, giving highly strategic, practical advice.\n\n"
        f"User's Resume:\n{resume_text}\n\n"
        f"Chat History:\n{formatted_history}\n"
        f"Real-Time Grounding Context:{grounding_str}\n"
        f"User's New Question: {message}\n"
        "Assistant: "
    )

    try:
        headers = {
            "Content-Type": "application/json",
            "Accept": "application/json",
            "Authorization": f"Bearer {mistral_key}"
        }
        payload = {
            "model": "mistral-large-latest",
            "messages": [{"role": "user", "content": prompt}],
            "stream": True
        }
        
        response = requests.post("https://api.mistral.ai/v1/chat/completions", headers=headers, json=payload, stream=True, timeout=30)
        if response.status_code != 200:
            raise ValueError(f"Mistral API returned status code {response.status_code}")

        for line in response.iter_lines():
            if line:
                decoded = line.decode('utf-8').strip()
                if decoded.startswith("data: "):
                    data_str = decoded[6:]
                    if data_str == "[DONE]":
                        break
                    try:
                        data_json = json.loads(data_str)
                        content = data_json['choices'][0]['delta'].get('content', '')
                        if content:
                            yield f"data: {json.dumps({'type': 'text', 'content': content})}\n\n"
                    except:
                        pass
                        
        yield f"data: {json.dumps({'type': 'done'})}\n\n"
    except Exception as e:
        yield f"data: {json.dumps({'type': 'error', 'message': f"Mistral API Error: {str(e)}"})}\n\n"

def build_improved_docx(text: str) -> str:
    """Builds a highly styled docx from the improved resume text and returns the file path."""
    doc_path = os.path.join(UPLOAD_DIR, "improved_resume.docx")
    
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)
        
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(51, 51, 51) # Charcoal
    
    lines = text.split("\n")
    for line in lines:
        cleaned = line.strip()
        if not cleaned:
            continue
            
        if cleaned.startswith("#") or (len(cleaned) < 40 and cleaned.isupper()):
            header_text = cleaned.replace("#", "").strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(header_text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(31, 78, 121) # Deep Blue
        elif cleaned.startswith("-") or cleaned.startswith("*"):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            bullet_text = cleaned[1:].strip()
            p.add_run(bullet_text)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.add_run(cleaned)
            
    doc.save(doc_path)
    return doc_path

def build_improved_pdf(text: str) -> str:
    """Builds a highly styled, executive-grade PDF from the improved resume text and returns the file path."""
    pdf_path = os.path.join(UPLOAD_DIR, "improved_resume.pdf")
    
    # Page setup: 0.75" margins = 54pt
    margin = 54
    doc = SimpleDocTemplate(
        pdf_path,
        pagesize=letter,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=margin,
        bottomMargin=margin
    )
    
    # Custom styles
    styles = getSampleStyleSheet()
    
    navy = colors.HexColor("#1F4E79")
    charcoal = colors.HexColor("#333333")
    steel = colors.HexColor("#4A6B82")
    
    name_style = ParagraphStyle(
        'ResumePDFName',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=navy,
        alignment=1, # Centered
        spaceAfter=4
    )
    
    contact_style = ParagraphStyle(
        'ResumePDFContact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=charcoal,
        alignment=1, # Centered
        spaceAfter=15
    )
    
    h1_style = ParagraphStyle(
        'ResumePDFH1',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=navy,
        spaceBefore=14,
        spaceAfter=4,
        keepWithNext=True
    )
    
    h2_style = ParagraphStyle(
        'ResumePDFH2',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=14,
        textColor=steel,
        spaceBefore=8,
        spaceAfter=3,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'ResumePDFBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=charcoal,
        spaceAfter=6
    )
    
    bullet_style = ParagraphStyle(
        'ResumePDFBullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=charcoal,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=4
    )

    story = []
    
    # Dynamic Premium Header (Shilpa's Contact Info)
    story.append(Paragraph("SHILPA NAIK, MHA", name_style))
    story.append(Paragraph("Orlando, FL | (407) 555-0199 | Shilpa.Naik2021@gmail.com | linkedin.com/in/shilpa-naik", contact_style))
    
    # Parse and structure lines
    lines = text.split("\n")
    processed_lines = []
    
    # 1. Skip duplicate contact info lines at the very top of the document
    in_resume_body = False
    for line in lines:
        cleaned = line.strip()
        if not cleaned:
            continue
            
        # Stop filtering once we hit the first main header
        if cleaned.startswith("#") or "SUMMARY" in cleaned.upper() or "EXPERIENCE" in cleaned.upper():
            in_resume_body = True
            
        if not in_resume_body:
            lower_line = cleaned.lower()
            # If it resembles header contact details, skip it to prevent duplication
            if (
                "shilpa" in lower_line or
                "naik" in lower_line or
                "orlando" in lower_line or
                "florida" in lower_line or
                "@" in lower_line or
                "linkedin" in lower_line or
                re.search(r'\d{3}[-\s]?\d{3}[-\s]?\d{4}', cleaned) or
                cleaned == "--" or
                cleaned == "• --"
            ):
                continue
                
        processed_lines.append(cleaned)

    # 2. Render each processed line with rich HTML styling
    for line in processed_lines:
        is_h1 = False
        is_h2 = False
        is_bullet = False
        
        cleaned_text = line
        
        # Categorize layout elements
        if cleaned_text.startswith("###"):
            is_h2 = True
            cleaned_text = cleaned_text.replace("###", "").strip()
        elif cleaned_text.startswith("##"):
            is_h2 = True
            cleaned_text = cleaned_text.replace("##", "").strip()
        elif cleaned_text.startswith("#"):
            is_h1 = True
            cleaned_text = cleaned_text.replace("#", "").strip()
        elif cleaned_text.startswith("-") or cleaned_text.startswith("*") or cleaned_text.startswith("•"):
            is_bullet = True
            cleaned_text = re.sub(r'^[-*•]\s*', '', cleaned_text).strip()
            
        # Clean section and sub-headers from raw markdown asterisks
        if is_h1 or is_h2:
            cleaned_text = cleaned_text.replace("**", "").replace("*", "").strip()
        else:
            # Escape XML special chars for safe ReportLab paragraph parsing
            cleaned_text = cleaned_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            # Convert markdown bold **text** -> HTML <b>text</b>
            cleaned_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', cleaned_text)
            # Convert markdown italic *text* -> HTML <i>text</i>
            cleaned_text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', cleaned_text)
            # Re-convert parsed XML characters inside HTML tags
            cleaned_text = cleaned_text.replace("&lt;b&gt;", "<b>").replace("&lt;/b&gt;", "</b>")
            cleaned_text = cleaned_text.replace("&lt;i&gt;", "<i>").replace("&lt;/i&gt;", "</i>")
            
            # Skip any blank separators or dashes
            if cleaned_text == "--" or cleaned_text == "• --" or not cleaned_text:
                continue

        # Render styled PDF flowables
        if is_h1:
            story.append(Paragraph(cleaned_text.upper(), h1_style))
            story.append(HRFlowable(width="100%", thickness=1, color=steel, spaceBefore=1, spaceAfter=8))
        elif is_h2:
            story.append(Paragraph(cleaned_text, h2_style))
        elif is_bullet:
            story.append(Paragraph(f"&bull;&nbsp; {cleaned_text}", bullet_style))
        else:
            story.append(Paragraph(cleaned_text, body_style))
            
    doc.build(story)
    return pdf_path
